# Gerador de Relatórios para Análise CSV
import pandas as pd
import numpy as np
from datetime import datetime
import io
import base64
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO

class ReportGenerator:
    """Classe para gerar relatórios em PDF e Word"""
    
    def __init__(self, analysis_name="Análise CSV", analysis_description=""):
        self.analysis_name = analysis_name
        self.analysis_description = analysis_description
        self.timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    def clean_text_for_pdf(self, text):
        """Remove emojis e caracteres especiais para compatibilidade com PDF"""
        if not text:
            return text
        
        # Mapeamento de emojis para texto
        emoji_map = {
            '👤': '[Usuario]',
            '🤖': '[IA]',
            '•': '-',
            '📊': '[Grafico]',
            '🔍': '[Analise]',
            '📈': '[Tendencia]',
            '⚠️': '[Atencao]',
            '💡': '[Dica]',
            '🎯': '[Objetivo]',
            '📋': '[Relatorio]',
            '🗑️': '[Lixeira]',
            '📤': '[Enviar]',
            '📥': '[Download]',
            '📄': '[Documento]',
            '📝': '[Texto]',
            '💬': '[Chat]',
            '🔧': '[Ferramenta]',
            '✅': '[OK]',
            '❌': '[Erro]',
            '⭐': '[Destaque]'
        }
        
        # Aplicar substituições
        clean_text = text
        for emoji, replacement in emoji_map.items():
            clean_text = clean_text.replace(emoji, replacement)
        
        # Remover caracteres problemáticos para PDF
        import re
        # Remover caracteres não-ASCII problemáticos
        clean_text = re.sub(r'[^\x00-\x7F]+', '?', clean_text)
        
        # Remover caracteres de controle problemáticos
        clean_text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', clean_text)
        
        # Remover quebras de linha excessivas
        clean_text = re.sub(r'\n+', ' ', clean_text)
        
        # Limitar tamanho do texto para evitar problemas de layout
        if len(clean_text) > 500:
            clean_text = clean_text[:500] + "..."
        
        # Garantir que o texto não esteja vazio
        if not clean_text.strip():
            clean_text = "Texto não disponível"
        
        return clean_text
    
    def generate_pdf_report(self, df, analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
        """Gera relatório em PDF com dados completos - versão ultra simplificada"""
        try:
            # Usar a função que sabemos que funciona
            return self._generate_simple_pdf(df, analysis_results, conversation_data, overview_data, crewai_conclusions)
        except Exception as e:
            raise Exception(f"Erro ao gerar PDF: {str(e)}")
    
    def _generate_simple_pdf(self, df, analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
        """Função interna para gerar PDF focado nas análises dos agentes CrewAI"""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Configurar fonte
        pdf.set_font('Arial', 'B', 16)
        
        # Título - usar apenas caracteres ASCII seguros
        safe_title = "Relatorio de Analise - Agentes CrewAI"
        pdf.cell(0, 10, safe_title, 0, 1, 'C')
        pdf.ln(5)
        
        # Data e hora
        pdf.set_font('Arial', '', 10)
        pdf.cell(0, 10, f"Gerado em: {self.timestamp}", 0, 1, 'C')
        pdf.ln(10)
        
        # Resumo dos dados básico
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, "Resumo dos Dados:", 0, 1)
        pdf.set_font('Arial', '', 10)
        
        pdf.cell(0, 6, f"Total de registros: {len(df)}", 0, 1)
        pdf.cell(0, 6, f"Total de colunas: {len(df.columns)}", 0, 1)
        pdf.ln(10)
        
        # Análises dos Agentes CrewAI
        if crewai_conclusions:
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, "ANALISES DOS AGENTES CREWAI", 0, 1, 'C')
            pdf.ln(5)
            
            # Mapear agentes para nomes em português
            agent_names = {
                "Data Validator": "Validador de Dados",
                "Data Profiler": "Perfilador de Dados", 
                "Pattern Detective": "Detetive de Padroes",
                "Anomaly Hunter": "Cacador de Anomalias",
                "Relationship Analyst": "Analista de Relacionamentos",
                "Strategic Synthesizer": "Sintetizador Estrategico"
            }
            
            for agent_name, agent_data in crewai_conclusions.items():
                # Nome do agente
                display_name = agent_names.get(agent_name, agent_name)
                pdf.set_font('Arial', 'B', 12)
                pdf.cell(0, 8, f"{display_name}:", 0, 1)
                
                # Status
                status = agent_data.get('status', 'unknown')
                pdf.set_font('Arial', '', 10)
                pdf.cell(0, 6, f"Status: {status}", 0, 1)
                
                # Resultado
                result = agent_data.get('result', 'Nenhum resultado disponivel')
                if isinstance(result, str):
                    # Limitar tamanho do resultado para PDF
                    if len(result) > 1000:
                        result = result[:1000] + "... (resultado truncado)"
                    
                    # Quebrar texto em linhas
                    lines = result.split('\n')
                    for line in lines[:20]:  # Máximo 20 linhas por agente
                        if len(line) > 80:
                            line = line[:80] + "..."
                        pdf.cell(0, 5, line, 0, 1)
                
                pdf.ln(5)
        else:
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, "Nenhuma analise CrewAI disponivel", 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.cell(0, 6, "Execute uma analise com os agentes CrewAI primeiro", 0, 1)
        
        # Salvar PDF
        pdf_output = pdf.output(dest='S')
        if isinstance(pdf_output, str):
            return pdf_output.encode('latin-1')
        elif isinstance(pdf_output, bytearray):
            return bytes(pdf_output)
        return pdf_output
    
    def generate_word_report(self, df, analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
        """Gera relatório em Word focado nas análises dos agentes CrewAI"""
        try:
            doc = Document()
            
            # Título
            title = doc.add_heading("Relatório de Análise - Agentes CrewAI", 0)
            title.alignment = 1  # Centralizado
            
            # Data e hora
            doc.add_paragraph(f"Gerado em: {self.timestamp}")
            doc.add_paragraph("")  # Linha em branco
            
            # Resumo dos dados básico
            doc.add_heading("Resumo dos Dados", level=1)
            doc.add_paragraph(f"Total de registros: {len(df):,}")
            doc.add_paragraph(f"Total de colunas: {len(df.columns)}")
            doc.add_paragraph("")  # Linha em branco
            
            # Análises dos Agentes CrewAI
            doc.add_heading("Análises dos Agentes CrewAI", level=1)
            
            if crewai_conclusions:
                # Mapear agentes para nomes em português
                agent_names = {
                    "Data Validator": "Validador de Dados",
                    "Data Profiler": "Perfilador de Dados", 
                    "Pattern Detective": "Detetive de Padrões",
                    "Anomaly Hunter": "Caçador de Anomalias",
                    "Relationship Analyst": "Analista de Relacionamentos",
                    "Strategic Synthesizer": "Sintetizador Estratégico"
                }
                
                for agent_name, agent_data in crewai_conclusions.items():
                    # Nome do agente
                    display_name = agent_names.get(agent_name, agent_name)
                    doc.add_heading(f"{display_name}", level=2)
                    
                    # Status
                    status = agent_data.get('status', 'unknown')
                    doc.add_paragraph(f"Status: {status}")
                    
                    # Resultado
                    result = agent_data.get('result', 'Nenhum resultado disponível')
                    if isinstance(result, str):
                        # Quebrar em parágrafos
                        lines = result.split('\n')
                        for line in lines:
                            if line.strip():
                                doc.add_paragraph(line.strip())
                    
                    doc.add_paragraph("")  # Linha em branco
            else:
                doc.add_paragraph("Nenhuma análise CrewAI disponível.")
                doc.add_paragraph("Execute uma análise com os agentes CrewAI primeiro.")
            
            
            # ===== SEÇÃO: CONCLUSÕES DOS AGENTES CREWAI =====
            if crewai_conclusions:
                doc.add_paragraph("")  # Linha em branco
                doc.add_heading("Conclusões dos Agentes CrewAI", level=1)
                
                # Mapear nomes dos agentes
                agent_names = {
                    "validation": "Data Validator",
                    "profiling": "Data Profiler", 
                    "patterns": "Pattern Detective",
                    "anomalies": "Anomaly Hunter",
                    "relationships": "Relationship Analyst",
                    "synthesis": "Strategic Synthesizer"
                }
                
                for agent_key, agent_data in crewai_conclusions.items():
                    agent_name = agent_names.get(agent_key, agent_key.replace('_', ' ').title())
                    status = agent_data.get('status', 'unknown')
                    result = agent_data.get('result', 'Nenhum resultado')
                    
                    # Adicionar cabeçalho do agente
                    p = doc.add_paragraph()
                    p.add_run(f"{agent_name} ({status}): ").bold = True
                    
                    # Adicionar resultado
                    if isinstance(result, dict):
                        # Se for dict, formatar melhor
                        for key, value in result.items():
                            doc.add_paragraph(f"• {key}: {value}")
                    else:
                        # Se for string, adicionar diretamente
                        doc.add_paragraph(str(result))
                    
                    doc.add_paragraph("")  # Linha em branco
            
            # ===== SEÇÃO: ANÁLISE E RESULTADOS =====
            if analysis_results:
                doc.add_paragraph("")  # Linha em branco
                doc.add_heading("Resultados da Análise", level=1)
                
                if isinstance(analysis_results, dict):
                    for key, value in analysis_results.items():
                        doc.add_paragraph(f"{key}: {value}")
                else:
                    doc.add_paragraph(str(analysis_results))
            
            # Salvar documento em bytes
            doc_buffer = BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            return doc_buffer.getvalue()
            
        except Exception as e:
            raise Exception(f"Erro ao gerar Word: {str(e)}")
    
    def create_download_button(self, file_data, filename, file_type):
        """Cria botão de download para Streamlit"""
        try:
            # Tratar diferentes tipos de dados
            if isinstance(file_data, str):
                # Se é string, converter para bytes
                file_bytes = file_data.encode('latin-1')
            elif isinstance(file_data, bytearray):
                # Se é bytearray, converter para bytes
                file_bytes = bytes(file_data)
            elif isinstance(file_data, bytes):
                # Se já é bytes, usar diretamente
                file_bytes = file_data
            else:
                # Tentar converter para string primeiro, depois para bytes
                file_bytes = str(file_data).encode('latin-1')
            
            b64 = base64.b64encode(file_bytes).decode()
            href = f'<a href="data:application/{file_type};base64,{b64}" download="{filename}">📥 Download {filename}</a>'
            return href
        except Exception as e:
            return f"Erro ao criar botão de download: {str(e)}"

    def generate_markdown_report(self, df, analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
        """Gera relatório em Markdown focado nas análises dos agentes CrewAI"""
        try:
            markdown_content = []
            
            # Título
            markdown_content.append("# Relatório de Análise - Agentes CrewAI")
            markdown_content.append("")
            markdown_content.append(f"**Gerado em:** {self.timestamp}")
            markdown_content.append("")
            
            # Resumo dos dados básico
            markdown_content.append("## Resumo dos Dados")
            markdown_content.append("")
            markdown_content.append(f"- **Total de registros:** {len(df):,}")
            markdown_content.append(f"- **Total de colunas:** {len(df.columns)}")
            markdown_content.append("")
            
            # Análises dos Agentes CrewAI
            markdown_content.append("## Análises dos Agentes CrewAI")
            markdown_content.append("")
            
            if crewai_conclusions:
                # Mapear agentes para nomes em português
                agent_names = {
                    "Data Validator": "Validador de Dados",
                    "Data Profiler": "Perfilador de Dados", 
                    "Pattern Detective": "Detetive de Padrões",
                    "Anomaly Hunter": "Caçador de Anomalias",
                    "Relationship Analyst": "Analista de Relacionamentos",
                    "Strategic Synthesizer": "Sintetizador Estratégico"
                }
                
                for agent_name, agent_data in crewai_conclusions.items():
                    # Nome do agente
                    display_name = agent_names.get(agent_name, agent_name)
                    markdown_content.append(f"### {display_name}")
                    markdown_content.append("")
                    
                    # Status
                    status = agent_data.get('status', 'unknown')
                    markdown_content.append(f"**Status:** {status}")
                    markdown_content.append("")
                    
                    # Resultado
                    result = agent_data.get('result', 'Nenhum resultado disponível')
                    if isinstance(result, str):
                        # Quebrar em linhas e formatar
                        lines = result.split('\n')
                        for line in lines:
                            if line.strip():
                                markdown_content.append(line.strip())
                    
                    markdown_content.append("")
                    markdown_content.append("---")
                    markdown_content.append("")
            else:
                markdown_content.append("Nenhuma análise CrewAI disponível.")
                markdown_content.append("")
                markdown_content.append("Execute uma análise com os agentes CrewAI primeiro.")
                markdown_content.append("")
            
            # Adicionar seção de conversação se disponível
            if conversation_data and conversation_data.get('messages'):
                markdown_content.append("## Conversação com IA")
                markdown_content.append("")
                markdown_content.append(f"**Total de mensagens:** {len(conversation_data['messages'])}")
                markdown_content.append("")
                
                for i, message in enumerate(conversation_data['messages'], 1):
                    # Timestamp
                    timestamp = message.get('timestamp', '')
                    if timestamp:
                        try:
                            dt = datetime.fromisoformat(timestamp.replace('Z', '+00:00'))
                            formatted_time = dt.strftime("%d/%m/%Y %H:%M:%S")
                        except:
                            formatted_time = timestamp
                    else:
                        formatted_time = "Horário não disponível"
                    
                    markdown_content.append(f"### Mensagem {i} - {formatted_time}")
                    markdown_content.append("")
                    
                    # Pergunta do usuário
                    user_msg = message.get('user_message', '')
                    if user_msg:
                        markdown_content.append("**👤 Usuário:**")
                        markdown_content.append("")
                        markdown_content.append(f"> {user_msg}")
                        markdown_content.append("")
                    
                    # Resposta da IA
                    ai_response = message.get('ai_response', '')
                    if ai_response:
                        markdown_content.append("**🤖 IA:**")
                        markdown_content.append("")
                        # Quebrar resposta em linhas para melhor formatação
                        lines = ai_response.split('\n')
                        for line in lines:
                            if line.strip():
                                markdown_content.append(line.strip())
                        markdown_content.append("")
                    
                    # Indicar se há gráfico
                    if message.get('has_chart', False):
                        markdown_content.append("*[Esta resposta incluiu um gráfico visual]*")
                        markdown_content.append("")
                    
                    markdown_content.append("---")
                    markdown_content.append("")
            
            # Converter para string
            markdown_text = '\n'.join(markdown_content)
            return markdown_text.encode('utf-8')
            
        except Exception as e:
            raise Exception(f"Erro ao gerar relatório Markdown: {str(e)}")

# Funções de conveniência
def generate_pdf_report(df, analysis_name="Análise CSV", analysis_description="", analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
    """Função de conveniência para gerar PDF com dados completos"""
    generator = ReportGenerator(analysis_name, analysis_description)
    return generator.generate_pdf_report(df, analysis_results, conversation_data, overview_data, crewai_conclusions)

def generate_markdown_report(df, analysis_name="Análise CSV", analysis_description="", analysis_results=None, conversation_data=None, overview_data=None, crewai_conclusions=None):
    """Função de conveniência para gerar Markdown com dados completos"""
    generator = ReportGenerator(analysis_name, analysis_description)
    return generator.generate_markdown_report(df, analysis_results, conversation_data, overview_data, crewai_conclusions)
